import streamlit as st
import pandas as pd
from io import StringIO, BytesIO
from predict import predict_demo
from front import render_html
from docx import Document
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import docx
from matplotlib.colors import to_hex, to_rgb
from pydocx import PyDocX
st.set_page_config(page_title="Vietnamese NER", layout="wide")
st.title("🔍 Vietnamese Named Entity Recognition (NER) Application")

# === Convert color name to hex ===
def name_to_hex(color_name):
    rgb = to_rgb(color_name)
    hex_code = to_hex(rgb)
    return hex_code[1:].upper()

# === Generate .docx with highlighted entities ===
def make_highlight_docx(tokens, labels):
    label_colors = {
        "PER": "lightcoral",
        "ORG": "lightblue",
        "LOC": "lightgreen",
        "MISC": "violet"
    }

    doc = Document()
    p = doc.add_paragraph()

    for token, label in zip(tokens, labels):
        run = p.add_run(token + " ")
        #Revire here https://www.cnblogs.com/alex-bn-lee/p/17732468.html
        if label != "O":
            entity = label.split("-")[-1]
            color_name = label_colors.get(entity, "gray")
            hex_color = name_to_hex(color_name)

            rPr = run._element.get_or_add_rPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), hex_color)
            rPr.append(shd)

    return doc

# === Tabs ===
tab1, tab2, tab3 = st.tabs(["📊 Data Analysis", "📈 Training Results", "🧪 Model Demo"])

# --- Tab 1: Data Analysis ---
with tab1:
    st.header("📊 Data Analysis")
    df = pd.DataFrame({
        "Entity Type": ["PER", "LOC", "ORG", "MISC"],
        "Count": [3200, 2500, 1800, 900]
    })
    st.bar_chart(df.set_index("Entity Type"))

# --- Tab 2: Training Results ---
with tab2:
    st.header("📈 Training Results")
    loss = [0.9, 0.7, 0.5, 0.35, 0.28]
    epoch = [1, 2, 3, 4, 5]
    df_loss = pd.DataFrame({"Epoch": epoch, "Loss": loss})
    st.line_chart(df_loss.set_index("Epoch"))

    st.subheader("Model Evaluation")
    df_eval = pd.DataFrame({
        "Version": ["v1", "v2", "v3"],
        "F1-score": [0.78, 0.83, 0.86],
        "Accuracy": [0.81, 0.85, 0.88]
    })
    st.dataframe(df_eval)

# --- Tab 3: Model Demo ---
with tab3:
    st.header("🧪 Vietnamese Named Entity Recognition")

    option = st.radio("Choose input method", ["✍️ Enter text", "📄 Upload .txt or .docx file"], horizontal=True)
    text = ""

    if option == "✍️ Enter text":
        text = st.text_input("Enter Vietnamese text:", "Nguyễn Văn A đang làm việc tại Hà Nội")

        if st.button("Analyze"):
            if not text.strip():
                st.warning("Please enter some text.")
            else:
                tokens, labels = predict_demo(text)

                st.subheader("🔍 Detected Entities")
                entities = [(tok, lab) for tok, lab in zip(tokens, labels) if lab != "O"]

                if entities:
                    for tok, lab in entities:
                        st.markdown(f"🔹 **{tok}** — *{lab}*")
                else:
                    st.info("No named entities detected.")

                st.subheader("📌 Highlighted Text")
                html_result = render_html(tokens, labels)
                st.markdown(html_result, unsafe_allow_html=True)

    else:
        uploaded_file = st.file_uploader("📄 Upload a .txt or .docx file", type=["txt", "docx"])

        if uploaded_file:
            if uploaded_file.type == "text/plain":
                stringio = StringIO(uploaded_file.getvalue().decode("utf-8"))
                text = stringio.read()
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                text = PyDocX.to_text(uploaded_file)

            if st.button("Analyze"):
                tokens, labels = predict_demo(text)

                entities = [(tok, lab) for tok, lab in zip(tokens, labels) if lab != "O"]

                st.subheader("🔍 Detected Entities")
                if entities:
                    for tok, lab in entities:
                        st.markdown(f"🔹 **{tok}** — *{lab}*")
                else:
                    st.info("No named entities detected.")

                if len(text) < 500:
                    st.subheader("📌 Highlighted Text")
                    html_result = render_html(tokens, labels)
                    st.markdown(html_result, unsafe_allow_html=True)
                else:
                    st.info("Text is too long. Please download the result file below.")

                doc = make_highlight_docx(tokens, labels)
                output = BytesIO()
                doc.save(output)
                output.seek(0)

                st.download_button(
                    label="📥 Download Result (DOCX)",
                    data=output,
                    file_name="ner_result.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
